from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from .utils import clean_text


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class DocumentText:
    path: str
    filename: str
    suffix: str
    title: str
    text: str
    pages: list[str]


def iter_files(path: Path) -> Iterable[Path]:
    if path.is_file():
        if path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path
        return
    for item in sorted(path.rglob("*")):
        if item.is_file() and item.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield item


def read_pdf(path: Path) -> tuple[str, list[str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("请先安装 pypdf：pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(clean_text(page.extract_text() or ""))
        except Exception:
            pages.append("")
    return clean_text("\n\n".join(pages)), pages


def read_docx(path: Path) -> tuple[str, list[str]]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("请先安装 python-docx：pip install python-docx") from exc

    doc = docx.Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = clean_text("\n".join(parts))
    return text, [text]


def read_plain(path: Path) -> tuple[str, list[str]]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="gbk", errors="ignore")
    text = clean_text(text)
    return text, [text]


def read_document(path: Path) -> DocumentText:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, pages = read_pdf(path)
    elif suffix == ".docx":
        text, pages = read_docx(path)
    elif suffix in {".txt", ".md"}:
        text, pages = read_plain(path)
    else:
        raise ValueError(f"不支持的文件类型：{path.suffix}")

    return DocumentText(
        path=str(path.resolve()),
        filename=path.name,
        suffix=suffix,
        title=path.stem,
        text=text,
        pages=pages,
    )
